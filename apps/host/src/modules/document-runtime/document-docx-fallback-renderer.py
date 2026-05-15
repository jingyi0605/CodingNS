#!/usr/bin/env python3
"""使用 python-docx 生成最小可用的真实 DOCX 文档。"""

from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


def main() -> int:
    if len(sys.argv) != 3:
        raise ValueError("usage: document-docx-fallback-renderer.py <payload.json> <output.docx>")

    payload_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])

    payload = json.loads(payload_path.read_text(encoding="utf-8"))
    document = build_document(payload)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return 0


def build_document(payload: dict) -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(11)

    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run(str(payload.get("title", "未命名文档")))
    title_run.bold = True
    title_run.font.size = Pt(20)

    summary = normalize_text(payload.get("summary"))
    if summary:
        summary_paragraph = document.add_paragraph()
        summary_paragraph.style = document.styles["Normal"]
        summary_run = summary_paragraph.add_run(summary)
        summary_run.italic = True

    meta = document.add_paragraph()
    meta.style = document.styles["Normal"]
    meta_run = meta.add_run(
        f"模板：{payload.get('templateKey', 'unknown')}@{payload.get('templateVersion', 'unknown')}"
    )
    meta_run.font.size = Pt(9)

    sections = payload.get("sections") or []
    for section_payload in sections:
        heading = normalize_text(section_payload.get("heading")) or "正文"
        body = normalize_text(section_payload.get("body")) or ""

        heading_paragraph = document.add_paragraph()
        heading_paragraph.style = document.styles["Heading 1"]
        heading_paragraph.add_run(heading)

        for block in split_blocks(body):
            body_paragraph = document.add_paragraph()
            body_paragraph.style = document.styles["Normal"]
            body_paragraph.add_run(block)

    references = payload.get("references") or []
    if references:
        references_heading = document.add_paragraph()
        references_heading.style = document.styles["Heading 1"]
        references_heading.add_run("引用来源")

        for reference in references:
            title_text = normalize_text(reference.get("title")) or "未命名来源"
            details = []
            source_ref = normalize_text(reference.get("sourceRef"))
            target_anchor_key = normalize_text(reference.get("targetAnchorKey"))
            quote_text = normalize_text(reference.get("quoteText"))

            if source_ref:
                details.append(f"来源：{source_ref}")
            if target_anchor_key:
                details.append(f"锚点：{target_anchor_key}")

            reference_paragraph = document.add_paragraph(style=document.styles["Normal"])
            reference_paragraph.add_run(title_text).bold = True
            if details:
                reference_paragraph.add_run(f"（{' | '.join(details)}）")

            if quote_text:
                quote_paragraph = document.add_paragraph(style=document.styles["Normal"])
                quote_paragraph.paragraph_format.left_indent = Pt(18)
                quote_paragraph.add_run(f"引文：{quote_text}")

    annotations = payload.get("annotations") or []
    if annotations:
        annotations_heading = document.add_paragraph()
        annotations_heading.style = document.styles["Heading 1"]
        annotations_heading.add_run("批注记录")

        for annotation in annotations:
            anchor_type = normalize_text(annotation.get("anchorType")) or "unknown"
            anchor_key = normalize_text(annotation.get("anchorKey")) or "unknown"
            body_text = normalize_text(annotation.get("body")) or ""
            status = normalize_text(annotation.get("status")) or "open"
            created_by = normalize_text(annotation.get("createdBy")) or "unknown"

            annotation_paragraph = document.add_paragraph(style=document.styles["Normal"])
            annotation_paragraph.add_run(f"[{status}] ").bold = True
            annotation_paragraph.add_run(f"{anchor_type}:{anchor_key} ")
            annotation_paragraph.add_run(body_text)
            annotation_paragraph.add_run(f"（{created_by}）").italic = True

    return document


def split_blocks(body: str) -> list[str]:
    blocks = [part.strip() for part in body.replace("\r\n", "\n").split("\n\n")]
    return [block for block in blocks if block]


def normalize_text(value: object) -> str | None:
    if value is None:
        return None

    text = str(value).strip()
    return text or None


if __name__ == "__main__":
    raise SystemExit(main())
